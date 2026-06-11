"""ATS-safe résumé export.

Applicant Tracking Systems parse plain, single-column documents well and choke
on tables, columns, text boxes, images, headers/footers, and fancy glyphs.
These builders emit exactly that: standard ALL-CAPS section headers, a single
column, comma-separated skills, and ASCII "- " bullets. Pure functions over a
plain dict so the output is deterministic and unit-tested; the view assembles
the dict from the StructuredProfile.

Export dict shape::

    {
        'full_name': str, 'headline': str, 'email': str, 'phone': str,
        'location': str, 'links': [str, ...], 'summary': str,
        'skills': [str | {'name': str}, ...],
        'experiences': [{'title','company','location','start','end',
                         'is_current': bool, 'bullets': [str, ...]}, ...],
        'educations': [{'degree','field_of_study','institution','end','gpa'}, ...],
        'projects': [{'name','description'}, ...],
    }
"""

from __future__ import annotations

import io
import re

SECTION_SUMMARY = 'SUMMARY'
SECTION_SKILLS = 'SKILLS'
SECTION_EXPERIENCE = 'EXPERIENCE'
SECTION_EDUCATION = 'EDUCATION'
SECTION_PROJECTS = 'PROJECTS'

_TAG_RE = re.compile(r'<[^>]+>')
_WS_RE = re.compile(r'[ \t]+')
# Glyphs ATS parsers commonly mangle → ASCII equivalents.
_GLYPHS = {
    '•': '-', '‣': '-', '●': '-', '·': '-',
    '–': '-', '—': '-', '‘': "'", '’': "'",
    '“': '"', '”': '"', ' ': ' ',
}


def _clean(text) -> str:
    text = _TAG_RE.sub(' ', str(text or ''))
    for bad, good in _GLYPHS.items():
        text = text.replace(bad, good)
    return _WS_RE.sub(' ', text).strip()


def _skill_names(skills) -> list[str]:
    out: list[str] = []
    for s in skills or []:
        name = _clean(s.get('name') if isinstance(s, dict) else s)
        if name and name not in out:
            out.append(name)
    return out


def _date_range(start, end, is_current=False) -> str:
    start = _clean(start)
    end = 'Present' if is_current else (_clean(end) or 'Present')
    return f'{start} - {end}' if start else (end if end != 'Present' else '')


def _experience_header(exp: dict) -> str:
    title_company = ' - '.join(p for p in (_clean(exp.get('title')), _clean(exp.get('company'))) if p)
    meta = ', '.join(m for m in (_clean(exp.get('location')),
                                 _date_range(exp.get('start'), exp.get('end'), exp.get('is_current'))) if m)
    if title_company and meta:
        return f'{title_company} ({meta})'
    return title_company or meta


def _education_header(edu: dict) -> str:
    degree = ' in '.join(p for p in (_clean(edu.get('degree')), _clean(edu.get('field_of_study'))) if p)
    head = ' - '.join(p for p in (degree, _clean(edu.get('institution'))) if p)
    end = _clean(edu.get('end'))
    return f'{head} ({end})' if head and end else (head or end)


def build_ats_resume(profile: dict) -> str:
    """Render an ATS-safe plain-text résumé. Always ends with a single newline."""
    lines: list[str] = [(_clean(profile.get('full_name')) or 'Your Name').upper()]

    contacts = [_clean(c) for c in (
        profile.get('email'), profile.get('phone'), profile.get('location'),
        *(profile.get('links') or []),
    ) if _clean(c)]
    if contacts:
        lines.append(' | '.join(contacts))

    headline = _clean(profile.get('headline'))
    if headline:
        lines.append(headline)

    def section(title: str) -> None:
        lines.extend(['', title])

    summary = _clean(profile.get('summary'))
    if summary:
        section(SECTION_SUMMARY)
        lines.append(summary)

    skills = _skill_names(profile.get('skills'))
    if skills:
        section(SECTION_SKILLS)
        lines.append(', '.join(skills))

    experiences = profile.get('experiences') or []
    if experiences:
        section(SECTION_EXPERIENCE)
        for i, exp in enumerate(experiences):
            if i:
                lines.append('')
            header = _experience_header(exp)
            if header:
                lines.append(header)
            for bullet in exp.get('bullets') or []:
                bullet = _clean(bullet)
                if bullet:
                    lines.append(f'- {bullet}')

    educations = profile.get('educations') or []
    if educations:
        section(SECTION_EDUCATION)
        for edu in educations:
            header = _education_header(edu)
            if header:
                lines.append(header)
            gpa = _clean(edu.get('gpa'))
            if gpa:
                lines.append(f'GPA: {gpa}')

    projects = profile.get('projects') or []
    if projects:
        section(SECTION_PROJECTS)
        for proj in projects:
            name = _clean(proj.get('name'))
            desc = _clean(proj.get('description'))
            line = f'{name} - {desc}' if name and desc else (name or desc)
            if line:
                lines.append(line)

    return '\n'.join(lines).strip() + '\n'


def build_ats_docx(profile: dict) -> bytes:
    """Render the same résumé as a minimal, single-column .docx (no tables).

    Built by reusing `build_ats_resume` so the two formats never diverge:
    ALL-CAPS section headers become bold paragraphs, everything else is a plain
    paragraph — the simplest structure an ATS reliably parses."""
    from docx import Document

    headers = {SECTION_SUMMARY, SECTION_SKILLS, SECTION_EXPERIENCE,
               SECTION_EDUCATION, SECTION_PROJECTS}
    document = Document()
    text_lines = build_ats_resume(profile).split('\n')
    name = text_lines[0] if text_lines else ''
    document.add_heading(name, level=0)
    for line in text_lines[1:]:
        if not line.strip():
            continue
        para = document.add_paragraph()
        run = para.add_run(line)
        if line in headers:
            run.bold = True
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
