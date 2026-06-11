import io

from resumes.ats_export import build_ats_docx, build_ats_resume

FULL_PROFILE = {
    'full_name': 'Jane Doe',
    'headline': 'Backend Engineer',
    'email': 'jane@example.com',
    'phone': '+1-555-0100',
    'location': 'Remote',
    'links': ['linkedin.com/in/jane'],
    'summary': 'Backend engineer with <b>8 years</b> building APIs.',
    'skills': [{'name': 'Python'}, {'name': 'Django'}, 'Python'],  # dup + plain
    'experiences': [{
        'title': 'Senior Backend Engineer', 'company': 'Acme',
        'location': 'Remote', 'start': '2022', 'end': '', 'is_current': True,
        'bullets': ['Built ingestion pipeline', '• Cut latency 40%'],
    }],
    'educations': [{
        'degree': 'B.S.', 'field_of_study': 'Computer Science',
        'institution': 'State University', 'end': '2018', 'gpa': '3.8',
    }],
    'projects': [{'name': 'JobBot', 'description': 'An OSS job tracker'}],
}


def test_build_ats_resume_has_standard_sections_and_clean_text():
    out = build_ats_resume(FULL_PROFILE)
    assert out.startswith('JANE DOE\n')
    assert 'jane@example.com | +1-555-0100 | Remote | linkedin.com/in/jane' in out
    for header in ('SUMMARY', 'SKILLS', 'EXPERIENCE', 'EDUCATION', 'PROJECTS'):
        assert f'\n{header}\n' in out
    # HTML stripped, skills de-duplicated, current role + ASCII bullets
    assert '<b>' not in out
    assert 'Python, Django' in out and out.count('Python') == 1
    assert 'Senior Backend Engineer - Acme (Remote, 2022 - Present)' in out
    assert '- Built ingestion pipeline' in out
    assert '•' not in out  # bullet glyph normalised to "- "
    assert 'B.S. in Computer Science - State University (2018)' in out
    assert out.endswith('\n')


def test_build_ats_resume_minimal_profile_degrades_gracefully():
    out = build_ats_resume({'skills': ['Go']})
    assert out.startswith('YOUR NAME')
    assert 'SKILLS\nGo' in out
    # no empty sections for absent data
    assert 'EXPERIENCE' not in out
    assert 'EDUCATION' not in out


def test_build_ats_docx_is_a_valid_single_column_document():
    from docx import Document

    payload = build_ats_docx(FULL_PROFILE)
    assert isinstance(payload, bytes) and len(payload) > 0

    document = Document(io.BytesIO(payload))
    text = '\n'.join(p.text for p in document.paragraphs)
    assert 'JANE DOE' in text
    assert 'SKILLS' in text
    # ATS-safe: no tables in the document
    assert document.tables == []
